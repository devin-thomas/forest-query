from __future__ import annotations

import csv
import sqlite3
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "submission"
SQL_APPENDIX_PATH = ROOT / "forest_query_analysis.sql"


@dataclass(frozen=True)
class CountryRow:
    country_name: str
    region: str
    value: float


def load_database() -> sqlite3.Connection:
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.executescript(
        """
        CREATE TABLE forest_area (
            country_code TEXT,
            country_name TEXT,
            year INTEGER,
            forest_area_sqkm REAL
        );

        CREATE TABLE land_area (
            country_code TEXT,
            country_name TEXT,
            year INTEGER,
            total_area_sq_mi REAL
        );

        CREATE TABLE regions (
            country_name TEXT,
            country_code TEXT,
            region TEXT,
            income_group TEXT
        );
        """
    )

    for filename, table_name in [
        ("forest_area.csv", "forest_area"),
        ("land_area.csv", "land_area"),
        ("regions.csv", "regions"),
    ]:
        with (ROOT / filename).open(newline="", encoding="utf-8-sig") as handle:
            reader = csv.reader(handle)
            next(reader)
            cur.executemany(f"INSERT INTO {table_name} VALUES (?,?,?,?)", reader)

    cur.executescript(
        """
        CREATE VIEW forestation AS
        SELECT
            fa.country_code,
            fa.country_name,
            fa.year,
            fa.forest_area_sqkm,
            la.total_area_sq_mi,
            r.region,
            r.income_group,
            (fa.forest_area_sqkm / (CAST(NULLIF(la.total_area_sq_mi, '') AS REAL) * 2.59)) * 100 AS percent_forest
        FROM forest_area AS fa
        JOIN land_area AS la
            ON fa.country_code = la.country_code
           AND fa.year = la.year
        JOIN regions AS r
            ON fa.country_code = r.country_code;
        """
    )
    return conn


def one(cur: sqlite3.Cursor, sql: str, params: tuple = ()) -> sqlite3.Row:
    row = cur.execute(sql, params).fetchone()
    if row is None:
        raise ValueError("Expected at least one row")
    return row


def many(cur: sqlite3.Cursor, sql: str, params: tuple = ()) -> list[sqlite3.Row]:
    return cur.execute(sql, params).fetchall()


def fmt_sqkm(value: float) -> str:
    return f"{value:,.2f} sq km"


def fmt_pct(value: float) -> str:
    return f"{value:.2f}%"


def load_sql_appendix() -> str:
    return SQL_APPENDIX_PATH.read_text(encoding="utf-8").strip()


def build_metrics(conn: sqlite3.Connection) -> dict:
    cur = conn.cursor()

    world_1990 = one(
        cur,
        "SELECT forest_area_sqkm FROM forestation WHERE country_name = 'World' AND year = 1990",
    )["forest_area_sqkm"]
    world_2016 = one(
        cur,
        "SELECT forest_area_sqkm FROM forestation WHERE country_name = 'World' AND year = 2016",
    )["forest_area_sqkm"]
    world_loss = world_1990 - world_2016
    world_loss_pct = (world_loss / world_1990) * 100

    closest_country = one(
        cur,
        """
        SELECT
            country_name,
            (CAST(total_area_sq_mi AS REAL) * 2.59) AS total_area_sqkm
        FROM forestation
        WHERE year = 2016
          AND country_name <> 'World'
        ORDER BY ABS((CAST(total_area_sq_mi AS REAL) * 2.59) - ?) ASC
        LIMIT 1
        """,
        (world_loss,),
    )

    regional_table = many(
        cur,
        """
        SELECT
            region,
            ROUND(
                SUM(CASE WHEN year = 1990 THEN forest_area_sqkm END)
                / SUM(CASE WHEN year = 1990 THEN CAST(total_area_sq_mi AS REAL) * 2.59 END) * 100,
                2
            ) AS pct_1990,
            ROUND(
                SUM(CASE WHEN year = 2016 THEN forest_area_sqkm END)
                / SUM(CASE WHEN year = 2016 THEN CAST(total_area_sq_mi AS REAL) * 2.59 END) * 100,
                2
            ) AS pct_2016
        FROM forestation
        GROUP BY region
        ORDER BY region
        """,
    )

    region_1990 = many(
        cur,
        """
        SELECT
            region,
            ROUND(SUM(forest_area_sqkm) / SUM(CAST(total_area_sq_mi AS REAL) * 2.59) * 100, 2) AS pct
        FROM forestation
        WHERE year = 1990
        GROUP BY region
        ORDER BY pct DESC
        """,
    )
    region_2016 = many(
        cur,
        """
        SELECT
            region,
            ROUND(SUM(forest_area_sqkm) / SUM(CAST(total_area_sq_mi AS REAL) * 2.59) * 100, 2) AS pct
        FROM forestation
        WHERE year = 2016
        GROUP BY region
        ORDER BY pct DESC
        """,
    )

    decreased_regions = many(
        cur,
        """
        WITH regional_1990 AS (
            SELECT
                region,
                SUM(forest_area_sqkm) / SUM(CAST(total_area_sq_mi AS REAL) * 2.59) * 100 AS pct_1990
            FROM forestation
            WHERE year = 1990
            GROUP BY region
        ),
        regional_2016 AS (
            SELECT
                region,
                SUM(forest_area_sqkm) / SUM(CAST(total_area_sq_mi AS REAL) * 2.59) * 100 AS pct_2016
            FROM forestation
            WHERE year = 2016
            GROUP BY region
        )
        SELECT
            regional_1990.region,
            ROUND(regional_1990.pct_1990, 2) AS pct_1990,
            ROUND(regional_2016.pct_2016, 2) AS pct_2016
        FROM regional_1990
        JOIN regional_2016
            ON regional_1990.region = regional_2016.region
        WHERE regional_2016.pct_2016 < regional_1990.pct_1990
          AND regional_1990.region <> 'World'
        ORDER BY regional_1990.pct_1990 - regional_2016.pct_2016 DESC
        """,
    )

    top_amount_decrease = many(
        cur,
        """
        SELECT
            f1990.country_name,
            f1990.region,
            ROUND(f2016.forest_area_sqkm - f1990.forest_area_sqkm, 2) AS change_sqkm
        FROM forestation AS f1990
        JOIN forestation AS f2016
            ON f1990.country_code = f2016.country_code
        WHERE f1990.year = 1990
          AND f2016.year = 2016
          AND f1990.country_name <> 'World'
        ORDER BY f1990.forest_area_sqkm - f2016.forest_area_sqkm DESC, f1990.country_name
        LIMIT 5
        """,
    )

    top_percent_decrease = many(
        cur,
        """
        SELECT
            f1990.country_name,
            f1990.region,
            ROUND(
                ((f2016.forest_area_sqkm - f1990.forest_area_sqkm) / f1990.forest_area_sqkm) * 100,
                2
            ) AS percent_change
        FROM forestation AS f1990
        JOIN forestation AS f2016
            ON f1990.country_code = f2016.country_code
        WHERE f1990.year = 1990
          AND f2016.year = 2016
          AND f1990.country_name <> 'World'
          AND f1990.forest_area_sqkm > 0
        ORDER BY ((f1990.forest_area_sqkm - f2016.forest_area_sqkm) / f1990.forest_area_sqkm) DESC,
                 f1990.country_name
        LIMIT 5
        """,
    )

    top_increase_amount = many(
        cur,
        """
        SELECT
            f1990.country_name,
            f1990.region,
            ROUND(f2016.forest_area_sqkm - f1990.forest_area_sqkm, 2) AS increase_sqkm
        FROM forestation AS f1990
        JOIN forestation AS f2016
            ON f1990.country_code = f2016.country_code
        WHERE f1990.year = 1990
          AND f2016.year = 2016
          AND f1990.country_name <> 'World'
        ORDER BY increase_sqkm DESC, f1990.country_name
        LIMIT 2
        """,
    )

    top_increase_percent = many(
        cur,
        """
        SELECT
            f1990.country_name,
            f1990.region,
            ROUND(
                ((f2016.forest_area_sqkm - f1990.forest_area_sqkm) / f1990.forest_area_sqkm) * 100,
                2
            ) AS pct_increase
        FROM forestation AS f1990
        JOIN forestation AS f2016
            ON f1990.country_code = f2016.country_code
        WHERE f1990.year = 1990
          AND f2016.year = 2016
          AND f1990.country_name <> 'World'
          AND f1990.forest_area_sqkm > 0
        ORDER BY pct_increase DESC, f1990.country_name
        LIMIT 1
        """,
    )

    quartile_counts = many(
        cur,
        """
        SELECT
            CASE
                WHEN percent_forest <= 25 THEN '1st quartile (0%-25%)'
                WHEN percent_forest <= 50 THEN '2nd quartile (25%-50%)'
                WHEN percent_forest <= 75 THEN '3rd quartile (50%-75%)'
                ELSE '4th quartile (>75%)'
            END AS quartile,
            COUNT(*) AS country_count
        FROM forestation
        WHERE year = 2016
          AND country_name <> 'World'
          AND percent_forest IS NOT NULL
        GROUP BY quartile
        ORDER BY country_count DESC, quartile
        """,
    )

    top_quartile_countries = many(
        cur,
        """
        SELECT
            country_name,
            region,
            ROUND(percent_forest, 2) AS percent_forest
        FROM forestation
        WHERE year = 2016
          AND country_name <> 'World'
          AND percent_forest > 75
        ORDER BY percent_forest DESC, country_name
        """,
    )

    countries_above_us = one(
        cur,
        """
        SELECT COUNT(*) AS country_count
        FROM forestation
        WHERE year = 2016
          AND country_name <> 'World'
          AND percent_forest > (
              SELECT percent_forest
              FROM forestation
              WHERE country_name = 'United States'
                AND year = 2016
          )
        """,
    )["country_count"]

    return {
        "world_1990": world_1990,
        "world_2016": world_2016,
        "world_loss": world_loss,
        "world_loss_pct": world_loss_pct,
        "closest_country": closest_country,
        "regional_table": regional_table,
        "region_1990": region_1990,
        "region_2016": region_2016,
        "decreased_regions": decreased_regions,
        "top_amount_decrease": top_amount_decrease,
        "top_percent_decrease": top_percent_decrease,
        "top_increase_amount": top_increase_amount,
        "top_increase_percent": top_increase_percent,
        "quartile_counts": quartile_counts,
        "top_quartile_countries": top_quartile_countries,
        "countries_above_us": countries_above_us,
    }


def build_markdown(metrics: dict) -> str:
    region_2016 = metrics["region_2016"]
    region_1990 = metrics["region_1990"]
    region_2016_map = {row["region"]: row["pct"] for row in region_2016}
    region_1990_map = {row["region"]: row["pct"] for row in region_1990}
    decreased = metrics["decreased_regions"]
    biggest_gain = metrics["top_increase_amount"][0]
    second_gain = metrics["top_increase_amount"][1]
    percent_gain = metrics["top_increase_percent"][0]
    amount_rows = metrics["top_amount_decrease"]
    percent_rows = metrics["top_percent_decrease"]
    quartile_rows = metrics["quartile_counts"]
    top_quartile = metrics["top_quartile_countries"]
    sql_appendix = load_sql_appendix()

    lines = [
        "# Report for ForestQuery into Global Deforestation, 1990 to 2016",
        "",
        "ForestQuery is on a mission to combat deforestation around the world and to raise awareness about how land-use change affects communities and ecosystems. This report uses the supplied World Bank forest-area, land-area, and regional datasets to identify the clearest patterns in global, regional, and country-level forestation between 1990 and 2016.",
        "",
        "## 1. Global Situation",
        "",
        f"In 1990, the world had {fmt_sqkm(metrics['world_1990'])} of forest area. By 2016, that figure had fallen to {fmt_sqkm(metrics['world_2016'])}, a net loss of {fmt_sqkm(metrics['world_loss'])} or {fmt_pct(metrics['world_loss_pct'])}.",
        "",
        f"The amount of forest area lost during this period is closest to the 2016 total land area of {metrics['closest_country']['country_name']}, which had {fmt_sqkm(metrics['closest_country']['total_area_sqkm'])} of land area. That comparison helps frame the decline in more concrete terms: the world effectively lost forest cover on a scale comparable to an entire large country.",
        "",
        "## 2. Regional Outlook",
        "",
        f"In 2016, {fmt_pct(region_2016_map['World'])} of the world's land area was forested. The highest regional forest share was in {region_2016[0]['region']} at {fmt_pct(region_2016[0]['pct'])}, while the lowest was in {region_2016[-1]['region']} at {fmt_pct(region_2016[-1]['pct'])}.",
        "",
        f"In 1990, the world forest share was slightly higher at {fmt_pct(region_1990_map['World'])}. The highest regional forest share was again {region_1990[0]['region']} at {fmt_pct(region_1990[0]['pct'])}, and the lowest was again {region_1990[-1]['region']} at {fmt_pct(region_1990[-1]['pct'])}.",
        "",
        "| Region | 1990 Forest Percentage | 2016 Forest Percentage |",
        "| --- | ---: | ---: |",
    ]

    for row in metrics["regional_table"]:
        lines.append(f"| {row['region']} | {fmt_pct(row['pct_1990'])} | {fmt_pct(row['pct_2016'])} |")

    lines.extend(
        [
            "",
            f"Only two non-world regions saw their forest share decline from 1990 to 2016: {decreased[0]['region']} fell from {fmt_pct(decreased[0]['pct_1990'])} to {fmt_pct(decreased[0]['pct_2016'])}, and {decreased[1]['region']} fell from {fmt_pct(decreased[1]['pct_1990'])} to {fmt_pct(decreased[1]['pct_2016'])}. Every other region increased slightly, but those gains were not large enough to offset the declines in these two regions, so the world total still fell from {fmt_pct(region_1990_map['World'])} to {fmt_pct(region_2016_map['World'])}.",
            "",
            "## 3. Country-Level Detail",
            "",
            "### Success Stories",
            "",
            f"The strongest positive outlier in the dataset is {biggest_gain['country_name']}, which increased its forest area by {fmt_sqkm(biggest_gain['increase_sqkm'])} between 1990 and 2016. The next-largest gain came from {second_gain['country_name']}, but its increase of {fmt_sqkm(second_gain['increase_sqkm'])} was far smaller. This suggests ForestQuery can learn from countries that have paired large land bases with sustained reforestation or conservation gains.",
            "",
            f"When the analysis shifts from absolute change to percent change, the leader is {percent_gain['country_name']}, which increased forest area by {fmt_pct(percent_gain['pct_increase'])}. This highlights that smaller countries can deliver dramatic relative gains even when their raw square-kilometer increase is not the largest in the world.",
            "",
            "### Largest Concerns",
            "",
            "The five countries with the largest absolute decreases in forest area between 1990 and 2016 are shown below.",
            "",
            "| Country | Region | Forest Area Decrease |",
            "| --- | --- | ---: |",
        ]
    )

    for row in amount_rows:
        lines.append(f"| {row['country_name']} | {row['region']} | {fmt_sqkm(abs(row['change_sqkm']))} |")

    lines.extend(
        [
            "",
            "The next table shows the five largest percent decreases in forest area over the same period.",
            "",
            "| Country | Region | Pct Forest Area Decrease |",
            "| --- | --- | ---: |",
        ]
    )

    for row in percent_rows:
        lines.append(f"| {row['country_name']} | {row['region']} | {fmt_pct(abs(row['percent_change']))} |")

    lines.extend(
        [
            "",
            f"Four of the top five countries by percent forest loss are in {percent_rows[1]['region']}: {percent_rows[1]['country_name']}, {percent_rows[2]['country_name']}, {percent_rows[3]['country_name']}, and {percent_rows[4]['country_name']}. The fifth is {percent_rows[0]['country_name']} in {percent_rows[0]['region']}.",
            "",
            f"{amount_rows[3]['country_name']} is the only country that appears in the top five for both absolute loss and percent loss. That combination makes it an especially urgent case for ForestQuery because the country is losing forest at scale and at a severe relative rate.",
            "",
            "### Quartiles",
            "",
            "| Quartile | Number of Countries |",
            "| --- | ---: |",
        ]
    )

    for row in quartile_rows:
        lines.append(f"| {row['quartile']} | {row['country_count']} |")

    lines.extend(
        [
            "",
            f"The largest number of countries in 2016 were in the {quartile_rows[0]['quartile']}.",
            "",
            f"There were {len(top_quartile)} countries in the top quartile with more than 75% of land designated as forest.",
            "",
            "| Country | Region | Pct Designated as Forest |",
            "| --- | --- | ---: |",
        ]
    )

    for row in top_quartile:
        lines.append(f"| {row['country_name']} | {row['region']} | {fmt_pct(row['percent_forest'])} |")

    lines.extend(
        [
            "",
            f"In 2016, {metrics['countries_above_us']} countries had a higher forestation percentage than the United States.",
            "",
            "## 4. Recommendations",
            "",
            f"ForestQuery should prioritize interventions and partnerships in {amount_rows[0]['country_name']}, {amount_rows[1]['country_name']}, {amount_rows[3]['country_name']}, and {amount_rows[4]['country_name']}. These countries represent either extremely large absolute losses, severe relative losses, or both.",
            "",
            f"The regional pattern also argues for sustained attention in {decreased[0]['region']} and {decreased[1]['region']}, the only two non-world regions that moved backward over the period. Communications, grantmaking, and local partner support should be concentrated where these regional declines overlap with the most at-risk countries.",
            "",
            f"At the same time, ForestQuery should study positive cases such as {biggest_gain['country_name']} and {percent_gain['country_name']} to understand which policy, restoration, and land-management strategies might be adapted elsewhere. Learning from successful reforestation stories is likely to improve the impact of future campaigns.",
            "",
            "## 5. Appendix: SQL Queries Used",
            "",
            "```sql",
            sql_appendix,
            "```",
        ]
    )

    return "\n".join(lines) + "\n"


def add_docx_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_docx(metrics: dict, output_path: Path) -> None:
    region_2016 = metrics["region_2016"]
    region_1990 = metrics["region_1990"]
    region_2016_map = {row["region"]: row["pct"] for row in region_2016}
    region_1990_map = {row["region"]: row["pct"] for row in region_1990}
    decreased = metrics["decreased_regions"]
    biggest_gain = metrics["top_increase_amount"][0]
    second_gain = metrics["top_increase_amount"][1]
    percent_gain = metrics["top_increase_percent"][0]
    amount_rows = metrics["top_amount_decrease"]
    percent_rows = metrics["top_percent_decrease"]
    quartile_rows = metrics["quartile_counts"]
    top_quartile = metrics["top_quartile_countries"]
    sql_appendix = load_sql_appendix()

    doc = Document()
    doc.add_heading("Report for ForestQuery into Global Deforestation, 1990 to 2016", 0)

    doc.add_paragraph(
        "ForestQuery is on a mission to combat deforestation around the world and to raise awareness "
        "about how land-use change affects communities and ecosystems. This report uses the supplied "
        "World Bank forest-area, land-area, and regional datasets to identify the clearest patterns "
        "in global, regional, and country-level forestation between 1990 and 2016."
    )

    doc.add_heading("1. Global Situation", level=1)
    doc.add_paragraph(
        f"In 1990, the world had {fmt_sqkm(metrics['world_1990'])} of forest area. By 2016, that figure "
        f"had fallen to {fmt_sqkm(metrics['world_2016'])}, a net loss of {fmt_sqkm(metrics['world_loss'])} "
        f"or {fmt_pct(metrics['world_loss_pct'])}."
    )
    doc.add_paragraph(
        f"The amount of forest area lost during this period is closest to the 2016 total land area of "
        f"{metrics['closest_country']['country_name']}, which had "
        f"{fmt_sqkm(metrics['closest_country']['total_area_sqkm'])} of land area."
    )

    doc.add_heading("2. Regional Outlook", level=1)
    doc.add_paragraph(
        f"In 2016, {fmt_pct(region_2016_map['World'])} of the world's land area was forested. The highest "
        f"regional forest share was in {region_2016[0]['region']} at {fmt_pct(region_2016[0]['pct'])}, "
        f"while the lowest was in {region_2016[-1]['region']} at {fmt_pct(region_2016[-1]['pct'])}."
    )
    doc.add_paragraph(
        f"In 1990, the world forest share was {fmt_pct(region_1990_map['World'])}. The highest regional "
        f"forest share was again {region_1990[0]['region']} at {fmt_pct(region_1990[0]['pct'])}, and the "
        f"lowest was again {region_1990[-1]['region']} at {fmt_pct(region_1990[-1]['pct'])}."
    )
    add_docx_table(
        doc,
        ["Region", "1990 Forest Percentage", "2016 Forest Percentage"],
        [
            [row["region"], fmt_pct(row["pct_1990"]), fmt_pct(row["pct_2016"])]
            for row in metrics["regional_table"]
        ],
    )
    doc.add_paragraph(
        f"Only two non-world regions saw their forest share decline from 1990 to 2016: "
        f"{decreased[0]['region']} fell from {fmt_pct(decreased[0]['pct_1990'])} to "
        f"{fmt_pct(decreased[0]['pct_2016'])}, and {decreased[1]['region']} fell from "
        f"{fmt_pct(decreased[1]['pct_1990'])} to {fmt_pct(decreased[1]['pct_2016'])}."
    )

    doc.add_heading("3. Country-Level Detail", level=1)
    doc.add_heading("Success Stories", level=2)
    doc.add_paragraph(
        f"The strongest positive outlier in the dataset is {biggest_gain['country_name']}, which "
        f"increased its forest area by {fmt_sqkm(biggest_gain['increase_sqkm'])} between 1990 and 2016. "
        f"The next-largest gain came from {second_gain['country_name']} at {fmt_sqkm(second_gain['increase_sqkm'])}."
    )
    doc.add_paragraph(
        f"When the analysis shifts from absolute change to percent change, the leader is "
        f"{percent_gain['country_name']}, which increased forest area by {fmt_pct(percent_gain['pct_increase'])}."
    )

    doc.add_heading("Largest Concerns", level=2)
    add_docx_table(
        doc,
        ["Country", "Region", "Forest Area Decrease"],
        [
            [row["country_name"], row["region"], fmt_sqkm(abs(row["change_sqkm"]))]
            for row in amount_rows
        ],
    )
    add_docx_table(
        doc,
        ["Country", "Region", "Pct Forest Area Decrease"],
        [
            [row["country_name"], row["region"], fmt_pct(abs(row["percent_change"]))]
            for row in percent_rows
        ],
    )
    doc.add_paragraph(
        f"Four of the top five countries by percent forest loss are in {percent_rows[1]['region']}: "
        f"{percent_rows[1]['country_name']}, {percent_rows[2]['country_name']}, "
        f"{percent_rows[3]['country_name']}, and {percent_rows[4]['country_name']}. The fifth is "
        f"{percent_rows[0]['country_name']} in {percent_rows[0]['region']}."
    )
    doc.add_paragraph(
        f"{amount_rows[3]['country_name']} is the only country that appears in the top five for both "
        f"absolute loss and percent loss."
    )

    doc.add_heading("Quartiles", level=2)
    add_docx_table(
        doc,
        ["Quartile", "Number of Countries"],
        [[row["quartile"], str(row["country_count"])] for row in quartile_rows],
    )
    doc.add_paragraph(
        f"The largest number of countries in 2016 were in the {quartile_rows[0]['quartile']}."
    )
    add_docx_table(
        doc,
        ["Country", "Region", "Pct Designated as Forest"],
        [
            [row["country_name"], row["region"], fmt_pct(row["percent_forest"])]
            for row in top_quartile
        ],
    )
    doc.add_paragraph(
        f"In 2016, {metrics['countries_above_us']} countries had a higher forestation percentage than the United States."
    )

    doc.add_heading("4. Recommendations", level=1)
    doc.add_paragraph(
        f"ForestQuery should prioritize intervention and partnership opportunities in "
        f"{amount_rows[0]['country_name']}, {amount_rows[1]['country_name']}, {amount_rows[3]['country_name']}, "
        f"and {amount_rows[4]['country_name']}. These countries combine large-scale forest loss with severe "
        f"local risk."
    )
    doc.add_paragraph(
        f"The regional pattern also points to sustained attention in {decreased[0]['region']} and "
        f"{decreased[1]['region']}, the only two non-world regions that lost forest share from 1990 to 2016."
    )
    doc.add_paragraph(
        f"ForestQuery should also study positive outliers such as {biggest_gain['country_name']} and "
        f"{percent_gain['country_name']} to identify conservation or restoration strategies worth adapting elsewhere."
    )

    doc.add_heading("5. Appendix: SQL Queries Used", level=1)
    sql_paragraph = doc.add_paragraph(style="No Spacing")
    sql_run = sql_paragraph.add_run(sql_appendix)
    sql_run.font.name = "Courier New"
    sql_run.font.size = Pt(8)

    doc.save(output_path)


def build_pdf(markdown_text: str, metrics: dict, output_path: Path) -> None:
    styles = getSampleStyleSheet()
    code_style = ParagraphStyle(
        "CodeBlock",
        parent=styles["Code"] if "Code" in styles else styles["BodyText"],
        fontName="Courier",
        fontSize=7,
        leading=8.5,
        leftIndent=10,
        rightIndent=10,
        spaceBefore=6,
        spaceAfter=6,
    )
    story = []
    region_2016 = metrics["region_2016"]
    region_1990 = metrics["region_1990"]
    region_2016_map = {row["region"]: row["pct"] for row in region_2016}
    region_1990_map = {row["region"]: row["pct"] for row in region_1990}
    decreased = metrics["decreased_regions"]
    biggest_gain = metrics["top_increase_amount"][0]
    second_gain = metrics["top_increase_amount"][1]
    percent_gain = metrics["top_increase_percent"][0]
    amount_rows = metrics["top_amount_decrease"]
    percent_rows = metrics["top_percent_decrease"]
    quartile_rows = metrics["quartile_counts"]
    top_quartile = metrics["top_quartile_countries"]
    sql_appendix = load_sql_appendix()

    story.append(Paragraph("Report for ForestQuery into Global Deforestation, 1990 to 2016", styles["Title"]))
    story.append(Spacer(1, 0.18 * inch))
    story.append(
        Paragraph(
            "ForestQuery is on a mission to combat deforestation around the world and to raise awareness "
            "about how land-use change affects communities and ecosystems. This report uses the supplied "
            "World Bank forest-area, land-area, and regional datasets to identify the clearest patterns "
            "in global, regional, and country-level forestation between 1990 and 2016.",
            styles["BodyText"],
        )
    )
    story.append(Spacer(1, 0.18 * inch))

    def add_body(text: str) -> None:
        story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 0.12 * inch))

    def add_pdf_table(headers: list[str], rows: list[list[str]]) -> None:
        table = Table([headers] + rows, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAD3")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.beige]),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("1. Global Situation", styles["Heading1"]))
    story.append(Spacer(1, 0.08 * inch))
    add_body(
        f"In 1990, the world had {fmt_sqkm(metrics['world_1990'])} of forest area. By 2016, that figure "
        f"had fallen to {fmt_sqkm(metrics['world_2016'])}, a net loss of {fmt_sqkm(metrics['world_loss'])} "
        f"or {fmt_pct(metrics['world_loss_pct'])}."
    )
    add_body(
        f"The amount of forest area lost during this period is closest to the 2016 total land area of "
        f"{metrics['closest_country']['country_name']}, which had "
        f"{fmt_sqkm(metrics['closest_country']['total_area_sqkm'])} of land area. That comparison helps "
        "frame the decline in more concrete terms: the world effectively lost forest cover on a scale "
        "comparable to an entire large country."
    )

    story.append(Paragraph("2. Regional Outlook", styles["Heading1"]))
    story.append(Spacer(1, 0.08 * inch))
    add_body(
        f"In 2016, {fmt_pct(region_2016_map['World'])} of the world's land area was forested. The highest "
        f"regional forest share was in {region_2016[0]['region']} at {fmt_pct(region_2016[0]['pct'])}, "
        f"while the lowest was in {region_2016[-1]['region']} at {fmt_pct(region_2016[-1]['pct'])}."
    )
    add_body(
        f"In 1990, the world forest share was slightly higher at {fmt_pct(region_1990_map['World'])}. "
        f"The highest regional forest share was again {region_1990[0]['region']} at "
        f"{fmt_pct(region_1990[0]['pct'])}, and the lowest was again {region_1990[-1]['region']} at "
        f"{fmt_pct(region_1990[-1]['pct'])}."
    )
    add_pdf_table(
        ["Region", "1990 Forest Percentage", "2016 Forest Percentage"],
        [
            [row["region"], fmt_pct(row["pct_1990"]), fmt_pct(row["pct_2016"])]
            for row in metrics["regional_table"]
        ],
    )
    add_body(
        f"Only two non-world regions saw their forest share decline from 1990 to 2016: "
        f"{decreased[0]['region']} fell from {fmt_pct(decreased[0]['pct_1990'])} to "
        f"{fmt_pct(decreased[0]['pct_2016'])}, and {decreased[1]['region']} fell from "
        f"{fmt_pct(decreased[1]['pct_1990'])} to {fmt_pct(decreased[1]['pct_2016'])}. Every other region "
        f"increased slightly, but those gains were not large enough to offset the declines in these two "
        f"regions, so the world total still fell from {fmt_pct(region_1990_map['World'])} to "
        f"{fmt_pct(region_2016_map['World'])}."
    )

    story.append(Paragraph("3. Country-Level Detail", styles["Heading1"]))
    story.append(Spacer(1, 0.08 * inch))
    story.append(Paragraph("Success Stories", styles["Heading2"]))
    story.append(Spacer(1, 0.06 * inch))
    add_body(
        f"The strongest positive outlier in the dataset is {biggest_gain['country_name']}, which increased "
        f"its forest area by {fmt_sqkm(biggest_gain['increase_sqkm'])} between 1990 and 2016. The next-largest "
        f"gain came from {second_gain['country_name']}, but its increase of {fmt_sqkm(second_gain['increase_sqkm'])} "
        "was far smaller. This suggests ForestQuery can learn from countries that have paired large land bases "
        "with sustained reforestation or conservation gains."
    )
    add_body(
        f"When the analysis shifts from absolute change to percent change, the leader is "
        f"{percent_gain['country_name']}, which increased forest area by {fmt_pct(percent_gain['pct_increase'])}. "
        "This highlights that smaller countries can deliver dramatic relative gains even when their raw "
        "square-kilometer increase is not the largest in the world."
    )
    story.append(Paragraph("Largest Concerns", styles["Heading2"]))
    story.append(Spacer(1, 0.06 * inch))
    add_body("The five countries with the largest absolute decreases in forest area between 1990 and 2016 are shown below.")
    add_pdf_table(
        ["Country", "Region", "Forest Area Decrease"],
        [
            [row["country_name"], row["region"], fmt_sqkm(abs(row["change_sqkm"]))]
            for row in metrics["top_amount_decrease"]
        ],
    )
    add_body("The next table shows the five largest percent decreases in forest area over the same period.")
    add_pdf_table(
        ["Country", "Region", "Pct Forest Area Decrease"],
        [
            [row["country_name"], row["region"], fmt_pct(abs(row["percent_change"]))]
            for row in metrics["top_percent_decrease"]
        ],
    )
    add_body(
        f"Four of the top five countries by percent forest loss are in {percent_rows[1]['region']}: "
        f"{percent_rows[1]['country_name']}, {percent_rows[2]['country_name']}, {percent_rows[3]['country_name']}, "
        f"and {percent_rows[4]['country_name']}. The fifth is {percent_rows[0]['country_name']} in "
        f"{percent_rows[0]['region']}."
    )
    add_body(
        f"{amount_rows[3]['country_name']} is the only country that appears in the top five for both absolute "
        "loss and percent loss. That combination makes it an especially urgent case for ForestQuery because "
        "the country is losing forest at scale and at a severe relative rate."
    )
    story.append(Paragraph("Quartiles", styles["Heading2"]))
    story.append(Spacer(1, 0.06 * inch))
    add_pdf_table(
        ["Quartile", "Number of Countries"],
        [[row["quartile"], str(row["country_count"])] for row in metrics["quartile_counts"]],
    )
    add_body(f"The largest number of countries in 2016 were in the {quartile_rows[0]['quartile']}.")
    add_body(f"There were {len(top_quartile)} countries in the top quartile with more than 75% of land designated as forest.")
    add_pdf_table(
        ["Country", "Region", "Pct Designated as Forest"],
        [
            [row["country_name"], row["region"], fmt_pct(row["percent_forest"])]
            for row in metrics["top_quartile_countries"]
        ],
    )
    add_body(
        f"In 2016, {metrics['countries_above_us']} countries had a higher forestation percentage than the United States."
    )

    story.append(Paragraph("4. Recommendations", styles["Heading1"]))
    story.append(Spacer(1, 0.08 * inch))
    add_body(
        f"ForestQuery should prioritize interventions and partnerships in {amount_rows[0]['country_name']}, "
        f"{amount_rows[1]['country_name']}, {amount_rows[3]['country_name']}, and {amount_rows[4]['country_name']}. "
        "These countries represent either extremely large absolute losses, severe relative losses, or both."
    )
    add_body(
        f"The regional pattern also argues for sustained attention in {decreased[0]['region']} and "
        f"{decreased[1]['region']}, the only two non-world regions that moved backward over the period. "
        "Communications, grantmaking, and local partner support should be concentrated where these regional "
        "declines overlap with the most at-risk countries."
    )
    add_body(
        f"At the same time, ForestQuery should study positive cases such as {biggest_gain['country_name']} and "
        f"{percent_gain['country_name']} to understand which policy, restoration, and land-management strategies "
        "might be adapted elsewhere. Learning from successful reforestation stories is likely to improve the impact "
        "of future campaigns."
    )

    story.append(Paragraph("5. Appendix: SQL Queries Used", styles["Heading1"]))
    story.append(Spacer(1, 0.08 * inch))
    story.append(Preformatted(sql_appendix, code_style))

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36,
    )
    document.build(story)


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    conn = load_database()
    metrics = build_metrics(conn)
    markdown_text = build_markdown(metrics)

    markdown_path = OUTPUT_DIR / "forest_query_report.md"
    docx_path = OUTPUT_DIR / "forest_query_report.docx"
    pdf_path = OUTPUT_DIR / "forest_query_report.pdf"

    markdown_path.write_text(markdown_text, encoding="utf-8")
    build_docx(metrics, docx_path)
    build_pdf(markdown_text, metrics, pdf_path)

    print(f"Wrote {markdown_path}")
    print(f"Wrote {docx_path}")
    print(f"Wrote {pdf_path}")


if __name__ == "__main__":
    main()
